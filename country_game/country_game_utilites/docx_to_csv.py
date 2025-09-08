import os
import csv
from datetime import datetime

try:
    from docx import Document  # python-docx
except Exception as e:
    Document = None

BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
COUNTRIES_DIR = os.path.join(BASE_DIR, 'countries')
OUTPUT_CSV = os.path.join(BASE_DIR, 'countries_docx_extract.csv')


def extract_text_from_docx(file_path: str) -> str:
    if not Document:
        raise RuntimeError('python-docx is not installed. Please install it with: pip install python-docx')
    doc = Document(file_path)
    parts = []
    for p in doc.paragraphs:
        text = (p.text or '').strip()
        if text:
            parts.append(text)
    # Include simple table text as well
    for tbl in doc.tables:
        for row in tbl.rows:
            cells = [('\n'.join(par.text for par in cell.paragraphs)).strip() for cell in row.cells]
            row_text = ' | '.join([c for c in cells if c])
            if row_text:
                parts.append(row_text)
    return '\n'.join(parts)


essential_fields = (
    'name', 'file', 'size_bytes', 'modified_iso', 'text'
)


def build_csv():
    if not os.path.isdir(COUNTRIES_DIR):
        raise FileNotFoundError(f"Countries directory not found: {COUNTRIES_DIR}")

    rows = []
    for fname in sorted(os.listdir(COUNTRIES_DIR)):
        if not fname.lower().endswith('.docx'):
            continue
        fpath = os.path.join(COUNTRIES_DIR, fname)
        try:
            text = extract_text_from_docx(fpath)
        except Exception as e:
            text = f"<ERROR: {e}>"
        stat = os.stat(fpath)
        name = os.path.splitext(fname)[0]
        rows.append({
            'name': name,
            'file': fpath,
            'size_bytes': stat.st_size,
            'modified_iso': datetime.fromtimestamp(stat.st_mtime).isoformat(timespec='seconds'),
            'text': text,
        })

    # Write CSV
    with open(OUTPUT_CSV, 'w', newline='', encoding='utf-8') as f:
        writer = csv.DictWriter(f, fieldnames=essential_fields)
        writer.writeheader()
        for r in rows:
            writer.writerow({k: r.get(k, '') for k in essential_fields})

    return OUTPUT_CSV, len(rows)


if __name__ == '__main__':
    out, n = build_csv()
    print(f"Wrote {n} records to {out}")
